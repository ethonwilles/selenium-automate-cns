import docx
import os
import requests


def invoice(loc="n/a",window_amount=0,slider_amt=0,amtofdoors=0,amtoflgdoors=0,amtoflgwin=0,desc_of_serv="n/a"):
    new_inv_num = requests.get("https://cns-automate-backend.herokuapp.com/invoice-tracker").json()["num"]
    try:
        inv_num = new_inv_num
        location = loc
        lot_num = 0
        window_amt = window_amount
        window_price_per = 20.5
        sliders = slider_amt
        slider_price_per = 38.5
        doors = amtofdoors
        doors_price_per = 35
        lgDoors = amtoflgdoors
        lgDoor_price_per = 20
        lgWin = amtoflgwin
        lgWin_price_per = 11
        price_of_trip = 0
        desc = desc_of_serv

        total = ((int(window_amt) * int(window_price_per)) +(int(lgDoors) * int(lgDoor_price_per)) + (int(sliders) * int(slider_price_per)) + (int(doors) * int(doors_price_per))+(int(lgWin) * int(lgWin_price_per)) +int(price_of_trip))


        


        document = docx.Document()

        document.add_heading(f'                        Invoice #{inv_num}', 0)

        p = document.add_paragraph("""
        Chad J Willes Construction
        371 west 500 south
        Lehi UT, 84043
        801-706-8523
        """)


        document.add_paragraph("""
        To:
        Alside Exterior Building Products
        915 West 2610 South 
        Salt Lake City Ut 84119
        Att: Andrew Germaine
        """)


        main_p = document.add_paragraph("")
        main_p.add_run('This invoice is for all the work needed to install windows to the Alside Standard').bold = True
        main_p.add_run(f"\n {location}")
        main_p.add_run(f"\nLot: {lot_num}")
        main_p.add_run(f"\nWindows                   @${window_price_per} per x{window_amt}")
        main_p.add_run(f"\nSliders                        @${slider_price_per} per x{sliders}")
        main_p.add_run(f"\nDoors                          @${doors_price_per} per x{doors}")
        main_p.add_run(f"\nLarge Windows       @${lgWin_price_per} per x{lgWin}")
        main_p.add_run(f"\nLarge Sliders            @${lgDoor_price_per} per x{lgDoors}")
        main_p.add_run(f"\nTrip Price: ${price_of_trip}")
        main_p.add_run(f"\n{desc}")
        main_p.add_run(f"\nTOTAL                                                                        ${total}").underline = True






        
        try:
            document.save('C:/Users/Ethon/Documents/Invoices/'+f'Invoice for Alside {inv_num} {location}.docx')
            requests.post("https://cns-automate-backend.herokuapp.com/invoice-tracker", json={"num": (new_inv_num + 1)})

            # message = MIMEMultipart()


            # mail_content = '''
            # Invoice For Windows
            # '''

            # message.attach(MIMEText(mail_content, 'plain'))
            # message['Subject'] = "Invoice"
            # attach_file_name = f'C:/Users/terra/Documents/Invoices/Invoice for Alside {inv_num} {location}.docx'
            # attach_file = open(attach_file_name, 'rb') # Open the file as binary mode
            # payload = MIMEBase('application', 'octate-stream')
            # payload.set_payload((attach_file).read())
            # encoders.encode_base64(payload) #encode the attachment
            # #add payload header with filename
            # payload.add_header('Content-Decomposition', 'attachment', filename='Invoice1.docx')
            # message.attach(payload)


            # r = requests.post(
		    #     "https://api.mailgun.net/v3/sandbox2a89baf534dc4c6a88ecc092c09a65e8.mailgun.org/messages.mime",
		    #     auth=("api", "9077b561bf5dbb09635c5372373e3182-f696beb4-2fb9e654"),
		    #     data={"from": "chadwilles@msn.com","to": "ethonwilles@gmail.com"},
            #     files={"message": bytes(str(message), "utf-8")})
            
        except Exception as e:
            print(f"\nInvoice Already Exists, Try Again. Error: {e}")
            
    except Exception as e:
        print(f"\n Program Broke for some reason. Try Again. \n Error Code for Ethon: {e}")
        

